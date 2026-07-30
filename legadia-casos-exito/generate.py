from __future__ import annotations

import html
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

BUILDER = "4.27.4"
ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
OUT.mkdir(parents=True, exist_ok=True)

COLORS = {
    "violet": "#261C30",
    "deep": "#031336",
    "blue": "#0B2B6B",
    "gold": "#E09900",
    "white": "#FFFFFF",
    "soft": "#F4F6FA",
    "muted": "#59647A",
}

URLS = {
    "home": "http://185.232.41.199/~wwwlegadianew/",
    "services": "http://185.232.41.199/~wwwlegadianew/servicios/",
    "cases": "http://185.232.41.199/~wwwlegadianew/casos-de-exito/",
    "contact": "http://185.232.41.199/~wwwlegadianew/contacto/",
    "phone": "tel:+34919359472",
    "email": "mailto:tecuidamos@legadia.es",
    "wa": "https://wa.me/34919359472?text=Hola%20Legadia%2C%20necesito%20ayuda%20con%20una%20herencia.",
}

IMAGES = {
    "logo": "http://185.232.41.199/~wwwlegadianew/wp-content/uploads/2025/12/logo-legadia-blanco.png",
    "hero": "https://images.unsplash.com/photo-1450101499163-c8848c66ca85?auto=format&fit=crop&w=1900&q=88",
    "madrid": "https://images.unsplash.com/photo-1543783207-ec64e4d95325?auto=format&fit=crop&w=1400&q=86",
    "avila": "https://images.unsplash.com/photo-1500534314209-a25ddb2bd429?auto=format&fit=crop&w=1400&q=86",
    "zamora": "https://images.unsplash.com/photo-1470770841072-f978cf4d019e?auto=format&fit=crop&w=1400&q=86",
    "soria": "https://images.unsplash.com/photo-1449157291145-7efd050a4d0e?auto=format&fit=crop&w=1400&q=86",
    "lugo": "https://images.unsplash.com/photo-1501854140801-50d01698950b?auto=format&fit=crop&w=1400&q=86",
    "jaen": "https://images.unsplash.com/photo-1470252649378-9c29740c9fa8?auto=format&fit=crop&w=1400&q=86",
    "asturias": "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?auto=format&fit=crop&w=1400&q=86",
    "rural": "https://images.unsplash.com/photo-1473445361085-b9a07f55608b?auto=format&fit=crop&w=1800&q=86",
    "asesoria": "https://images.unsplash.com/photo-1554224155-8d04cb21cd6c?auto=format&fit=crop&w=1200&q=86",
    "notaria": "https://images.unsplash.com/photo-1589829545856-d10d557cf95f?auto=format&fit=crop&w=1200&q=86",
    "inmobiliaria": "https://images.unsplash.com/photo-1560518883-ce09059eeffa?auto=format&fit=crop&w=1200&q=86",
    "cta": "https://images.unsplash.com/photo-1521737711867-e3b97375f902?auto=format&fit=crop&w=1800&q=86",
}

CASES = [
    ("Madrid capital", "Vivienda heredada y dos ramas familiares fuera de España", "Se reconstruyó la línea sucesoria, se ordenaron certificados y se localizó a familiares residentes en el extranjero.", "Personas localizadas y siguiente fase notarial definida.", IMAGES["madrid"]),
    ("Ávila rural", "Casa familiar y pequeñas parcelas sin adjudicar", "La revisión registral permitió separar inmuebles, cargas y titulares antes de iniciar nuevas gestiones.", "Inventario útil y ruta de tramitación acordada.", IMAGES["avila"]),
    ("Zamora y Argentina", "Una rama familiar emigrada que nadie conseguía contactar", "El contraste documental y genealógico conectó registros españoles con la rama familiar residente en Argentina.", "Contacto válido y expediente reactivado.", IMAGES["zamora"]),
    ("Soria", "Inmueble en un pueblo con sucesión incompleta", "Se priorizaron las búsquedas que podían cambiar el expediente y se descartaron gestiones repetidas.", "Rama sucesoria completada y documentación esencial reunida.", IMAGES["soria"]),
    ("Lugo", "Registros parroquiales, nombres repetidos y varias generaciones", "La investigación permitió distinguir personas, fechas y parentescos sin mezclar ramas familiares.", "Árbol familiar contrastado y herederos identificados.", IMAGES["lugo"]),
    ("Jaén", "Olivares, copropiedad y decisiones bloqueadas", "Se ordenaron titulares, porcentajes y documentos para valorar una tramitación o una salida económica.", "Derechos delimitados y alternativas explicadas.", IMAGES["jaen"]),
    ("Asturias y León", "Dos ramas familiares que habían perdido el contacto", "La coordinación documental evitó duplicidades y permitió explicar el caso a todos los interesados.", "Interlocutores localizados y plan común de actuación.", IMAGES["asturias"]),
]

VIDEOS = [
    ("Colaboración con una asesoría", "Certificados, impuestos y documentación coordinados sin peticiones contradictorias.", IMAGES["asesoria"]),
    ("Colaboración con un notario", "Información contrastada para que una declaración de herederos o adjudicación pueda avanzar.", IMAGES["notaria"]),
    ("Colaboración con una inmobiliaria", "Cómo desbloquear la comercialización cuando la titularidad hereditaria no está ordenada.", IMAGES["inmobiliaria"]),
]


def esc(value: object) -> str:
    return html.escape(str(value), quote=True)


def attrs(values: dict[str, object]) -> str:
    return " ".join(f'{key}="{esc(value)}"' for key, value in values.items() if value is not None)


def sc(name: str, values: dict[str, object] | None = None, content: str = "") -> str:
    return f"[{name} {attrs(values or {})}]{content}[/{name}]"


def common(module_class: str) -> dict[str, object]:
    return {
        "_builder_version": BUILDER,
        "_module_preset": "default",
        "global_colors_info": "{}",
        "module_class": module_class,
    }


def section(content: str, label: str, module_class: str, **extra: object) -> str:
    values = common(module_class)
    values.update({"fb_built": "1", "admin_label": label})
    values.update(extra)
    return sc("et_pb_section", values, content)


def row(content: str, structure: str, module_class: str, **extra: object) -> str:
    values = common(module_class)
    values.update({
        "column_structure": structure,
        "width": "86%",
        "max_width": "1240px",
        "width_tablet": "92%",
        "width_phone": "calc(100% - 30px)",
        "custom_padding": "0px||0px||false|false",
    })
    values.update(extra)
    return sc("et_pb_row", values, content)


def col(content: str, ctype: str, module_class: str, **extra: object) -> str:
    values = common(module_class)
    values.update({"type": ctype})
    values.update(extra)
    return sc("et_pb_column", values, content)


def blurb(title: str, body: str, module_class: str, dark: bool = False, level: str = "h3", size: str = "29px", **extra: object) -> str:
    values = common(module_class)
    values.update({
        "title": title,
        "header_level": level,
        "use_icon": "off",
        "header_font": "Josefin Sans|300|||||||",
        "header_text_color": COLORS["white"] if dark else COLORS["blue"],
        "header_font_size": size,
        "header_font_size_tablet": "26px",
        "header_font_size_phone": "23px",
        "header_line_height": "1.18em",
        "body_font": "Josefin Sans|300|||||||",
        "body_text_color": "rgba(255,255,255,0.78)" if dark else COLORS["muted"],
        "body_font_size": "17px",
        "body_font_size_phone": "16px",
        "body_line_height": "1.62em",
        "animation_style": "fade",
    })
    values.update(extra)
    return sc("et_pb_blurb", values, esc(body) if body else "")


def heading(kicker: str, title: str, body: str, module_class: str, dark: bool = False) -> str:
    return blurb(kicker, "", module_class + "-kicker", dark=False, level="h4", size="14px", header_font="Catamaran|600||on|||||", header_text_color=COLORS["gold"], header_letter_spacing="2px", animation_style="none") + blurb(title, body, module_class, dark=dark, level="h2", size="48px", body_font_size="19px")


def button(text: str, url: str, module_class: str, outline: bool = False) -> str:
    values = common(module_class)
    values.update({
        "button_text": text,
        "button_url": url,
        "button_alignment": "left",
        "custom_button": "on",
        "button_text_size": "15px",
        "button_text_color": COLORS["white"],
        "button_bg_color": "RGBA(255,255,255,0)" if outline else COLORS["gold"],
        "button_border_width": "1px",
        "button_border_color": COLORS["white"] if outline else COLORS["gold"],
        "button_border_radius": "0px",
        "button_font": "Josefin Sans|600||on|||||",
        "custom_padding": "15px|28px|15px|28px|false|false",
        "custom_margin": "18px|10px|0px|0px|false|false",
        "custom_css_main_element": "display:inline-block!important;width:auto!important;",
        "custom_css_main_element_phone": "display:block!important;width:100%!important;",
        "animation_style": "none",
    })
    return sc("et_pb_button", values)


def image(url: str, alt: str, module_class: str, height: str = "520px") -> str:
    values = common(module_class)
    values.update({"src": url, "alt": alt, "force_fullwidth": "on", "height": height, "height_phone": "300px", "object_fit": "cover", "animation_style": "fade"})
    return sc("et_pb_image", values)


def counter(number: str, title: str, module_class: str) -> str:
    values = common(module_class)
    values.update({
        "number": number,
        "title": title,
        "percent_sign": "off",
        "number_font": "Poppins|300|||||||",
        "number_text_color": COLORS["gold"],
        "number_font_size": "48px",
        "title_font": "Josefin Sans|300|||||||",
        "title_text_color": COLORS["white"],
        "title_font_size": "16px",
        "text_orientation": "center",
    })
    return sc("et_pb_number_counter", values)


def slide(case: tuple[str, str, str, str, str], index: int) -> str:
    place, title, body, result, image_url = case
    values = common(f"legadia-cases-slide legadia-cases-slide-{index}")
    values.update({
        "heading": f"{place} · {title}",
        "button_text": "Cuéntanos tu caso",
        "button_link": URLS["contact"],
        "background_image": image_url,
        "background_color": COLORS["deep"],
        "use_background_color_gradient": "on",
        "background_color_gradient_stops": "rgba(3,19,54,0.95) 0%|rgba(8,0,124,0.80) 60%|rgba(224,153,0,0.34) 120%",
        "background_color_gradient_overlays_image": "on",
        "header_font": "Josefin Sans|300|||||||",
        "header_text_color": COLORS["white"],
        "header_font_size": "38px",
        "header_font_size_phone": "28px",
        "body_font": "Josefin Sans|300|||||||",
        "body_text_color": "rgba(255,255,255,0.84)",
        "body_font_size": "18px",
        "button_text_color": COLORS["white"],
        "button_bg_color": COLORS["gold"],
        "button_border_color": COLORS["gold"],
        "button_border_radius": "0px",
        "custom_padding": "72px|8%|72px|8%|false|false",
    })
    return sc("et_pb_slide", values, esc(body + " Resultado: " + result))


def video(title: str, body: str, poster: str, index: int) -> str:
    module = sc("et_pb_video", {**common(f"legadia-cases-video-{index}"), "src": "", "image_src": poster, "play_icon_color": COLORS["gold"], "height": "310px", "height_phone": "240px", "admin_label": title})
    return module + blurb(title, body, f"legadia-cases-video-copy-{index}", size="25px", header_font="Josefin Sans|600|||||||")


def testimonial(author: str, body: str, index: int) -> str:
    values = common(f"legadia-cases-testimonial-{index}")
    values.update({
        "author": author,
        "job_title": "Caso anonimizado",
        "company_name": "Particular",
        "quote_icon_color": COLORS["gold"],
        "body_font": "Josefin Sans|300|||||||",
        "body_text_color": COLORS["muted"],
        "body_font_size": "17px",
        "author_font": "Josefin Sans|600|||||||",
        "author_text_color": COLORS["blue"],
        "background_color": COLORS["white"],
        "custom_padding": "30px|30px|30px|30px|false|false",
        "border_width_top": "3px",
        "border_color_top": COLORS["gold"],
    })
    return sc("et_pb_testimonial", values, esc(body))


def contact_form() -> str:
    fields = [
        sc("et_pb_contact_field", {**common("legadia-cases-field-name"), "field_id": "Name", "field_title": "Nombre y apellidos", "field_type": "input", "required_mark": "required"}),
        sc("et_pb_contact_field", {**common("legadia-cases-field-phone"), "field_id": "Phone", "field_title": "Teléfono", "field_type": "input", "required_mark": "required"}),
        sc("et_pb_contact_field", {**common("legadia-cases-field-email"), "field_id": "Email", "field_title": "Email", "field_type": "email", "fullwidth_field": "on", "required_mark": "required"}),
        sc("et_pb_contact_field", {**common("legadia-cases-field-place"), "field_id": "Place", "field_title": "Provincia o localidad del caso", "field_type": "input", "fullwidth_field": "on"}),
        sc("et_pb_contact_field", {**common("legadia-cases-field-type"), "field_id": "Type", "field_title": "Qué necesitas resolver", "field_type": "select", "field_options": "Localizar herederos|Ordenar documentación|Tramitar la herencia|Valorar derechos hereditarios|Coordinar un inmueble|No lo tengo claro", "fullwidth_field": "on", "required_mark": "required"}),
        sc("et_pb_contact_field", {**common("legadia-cases-field-message"), "field_id": "Message", "field_title": "Cuéntanos brevemente qué ocurre", "field_type": "text", "fullwidth_field": "on", "required_mark": "required"}),
    ]
    values = common("legadia-cases-contact-form")
    values.update({
        "email": "tecuidamos@legadia.es",
        "success_message": "Hemos recibido tu consulta.",
        "submit_button_text": "Enviar consulta",
        "custom_button": "on",
        "button_text_color": COLORS["white"],
        "button_bg_color": COLORS["gold"],
        "button_border_color": COLORS["gold"],
        "button_border_radius": "0px",
        "button_font": "Josefin Sans|600||on|||||",
        "form_field_font": "Josefin Sans|300|||||||",
        "form_field_text_color": COLORS["blue"],
        "form_field_background_color": COLORS["white"],
        "form_field_border_width_bottom": "2px",
        "form_field_border_color_bottom": COLORS["gold"],
        "custom_padding": "34px|34px|34px|34px|false|false",
        "background_color": COLORS["white"],
    })
    return sc("et_pb_contact_form", values, "".join(fields))


def build_layout() -> str:
    header = section(
        row(
            col(image(IMAGES["logo"], "Legadia", "legadia-cases-header-logo", "74px"), "1_4", "legadia-cases-header-logo-col")
            + col(button("Servicios", URLS["services"], "legadia-cases-header-services", True), "1_4", "legadia-cases-header-nav-col-1")
            + col(button("Casos de éxito", URLS["cases"], "legadia-cases-header-cases", True), "1_4", "legadia-cases-header-nav-col-2")
            + col(button("Contacto", URLS["contact"], "legadia-cases-header-contact"), "1_4", "legadia-cases-header-nav-col-3"),
            "1_4,1_4,1_4,1_4", "legadia-cases-header-row", make_equal="on"),
        "00 · Cabecera interna", "legadia-cases-section-header", background_color=COLORS["blue"], custom_padding="14px||14px||true|false")

    hero_copy = heading("Casos particulares", "Herencias que volvieron a avanzar", "Resultados explicados con casos anonimizados de Madrid, capitales de provincia y España rural.", "legadia-cases-hero-heading", True)
    hero_copy += blurb("Madrid y toda España", "Personas localizadas, documentación ordenada y expedientes que dejaron de estar paralizados.", "legadia-cases-hero-note", True, size="22px", background_color="rgba(255,255,255,0.08)", custom_padding="20px|22px|20px|22px|false|false", border_width_all="1px", border_color_all="rgba(255,255,255,0.24)")
    hero_copy += button("Ver los casos", "#casos-particulares", "legadia-cases-hero-primary") + button("Escribir WhatsApp", URLS["wa"], "legadia-cases-hero-secondary", True)
    hero = section(row(col(hero_copy, "3_5", "legadia-cases-hero-copy-col") + col(image(IMAGES["madrid"], "Caso de éxito en Madrid", "legadia-cases-hero-image"), "2_5", "legadia-cases-hero-image-col"), "3_5,2_5", "legadia-cases-hero-row", make_equal="on", gutter_width="2"), "01 · Hero enfocado", "legadia-cases-section-hero", background_color=COLORS["deep"], background_image=IMAGES["hero"], background_attachment="fixed", use_background_color_gradient="on", background_color_gradient_stops="rgba(8,0,124,0.95) 0%|rgba(3,19,54,0.92) 62%|rgba(224,153,0,0.48) 125%", background_color_gradient_overlays_image="on", custom_padding="120px||96px||true|false")

    metrics = section(row(
        col(counter("+400", "herederos encontrados", "legadia-cases-counter-1"), "1_4", "legadia-cases-counter-col-1")
        + col(counter("7", "tipologías de casos", "legadia-cases-counter-2"), "1_4", "legadia-cases-counter-col-2")
        + col(counter("3", "colaboraciones profesionales", "legadia-cases-counter-3"), "1_4", "legadia-cases-counter-col-3")
        + col(counter("1", "siguiente paso claro", "legadia-cases-counter-4"), "1_4", "legadia-cases-counter-col-4"),
        "1_4,1_4,1_4,1_4", "legadia-cases-metrics-row", make_equal="on"), "02 · Resultados", "legadia-cases-section-metrics", background_color=COLORS["blue"], custom_padding="42px||42px||true|false")

    featured_copy = heading("Caso destacado · Madrid", "Una vivienda, seis herederos y dos ramas fuera de España", "El expediente no avanzaba porque faltaban certificados y cada persona había recibido una versión distinta.", "legadia-cases-featured-heading")
    featured_copy += blurb("El bloqueo", "No existía un interlocutor común y dos ramas familiares no estaban localizadas.", "legadia-cases-featured-block", size="22px")
    featured_copy += blurb("La intervención", "Se reconstruyó la línea sucesoria y se contrastaron datos civiles y registrales.", "legadia-cases-featured-intervention", size="22px")
    featured_copy += blurb("El resultado", "Las personas quedaron identificadas y pudo definirse la siguiente fase notarial.", "legadia-cases-featured-result", size="22px", header_text_color=COLORS["gold"])
    featured = section(row(col(image(IMAGES["madrid"], "Caso destacado Madrid", "legadia-cases-featured-image", "610px"), "1_2", "legadia-cases-featured-image-col") + col(featured_copy, "1_2", "legadia-cases-featured-copy-col", background_color=COLORS["white"], custom_padding="54px|48px|54px|48px|false|false"), "1_2,1_2", "legadia-cases-featured-row", make_equal="on", gutter_width="1", width="92%", max_width="1320px"), "03 · Caso destacado", "legadia-cases-section-featured", background_color=COLORS["soft"], custom_padding="88px||88px||true|false")

    slider_values = common("legadia-cases-slider")
    slider_values.update({"show_arrows": "on", "show_pagination": "on", "auto": "off", "admin_label": "Carrusel de casos particulares"})
    slider = sc("et_pb_slider", slider_values, "".join(slide(case, i + 1) for i, case in enumerate(CASES)))
    carousel = section(
        row(col(heading("Casos particulares", "Madrid, provincias y España profunda", "Siete situaciones diferentes reunidas en un carrusel para mantener el foco en resultados.", "legadia-cases-carousel-heading", True), "1_1", "legadia-cases-carousel-head-col"), "1_1", "legadia-cases-carousel-head-row")
        + row(col(slider, "1_1", "legadia-cases-slider-col"), "1_1", "legadia-cases-slider-row"),
        "04 · Carrusel", "legadia-cases-section-carousel", background_color=COLORS["violet"], background_image=IMAGES["rural"], background_attachment="fixed", use_background_color_gradient="on", background_color_gradient_stops="rgba(38,28,48,0.96) 0%|rgba(8,0,124,0.91) 62%|rgba(11,43,107,0.90) 100%", background_color_gradient_overlays_image="on", custom_padding="96px||96px||true|false", custom_css_id="casos-particulares")

    rural_copy = heading("Cobertura real", "Del centro de Madrid a la España rural", "Registros antiguos, propiedades de varias generaciones, emigración y familias que dejaron de tener contacto.", "legadia-cases-rural-heading")
    rural_copy += blurb("Reconstruir", "Unir nombres, fechas, parentescos y domicilios sin confundir personas.", "legadia-cases-rural-1", size="22px")
    rural_copy += blurb("Verificar", "Contrastar información familiar con certificados, registros y bienes.", "legadia-cases-rural-2", size="22px")
    rural_copy += blurb("Contactar", "Explicar el motivo de la comunicación sin generar desconfianza.", "legadia-cases-rural-3", size="22px")
    rural = section(row(col(rural_copy, "1_2", "legadia-cases-rural-copy-col") + col(image(IMAGES["rural"], "Herencias en la España rural", "legadia-cases-rural-image", "560px"), "1_2", "legadia-cases-rural-image-col"), "1_2,1_2", "legadia-cases-rural-row", make_equal="on"), "05 · Madrid y España rural", "legadia-cases-section-rural", background_color=COLORS["white"], custom_padding="96px||96px||true|false")

    video_cols = "".join(col(video(title, body, poster, i), "1_3", f"legadia-cases-video-col-{i}") for i, (title, body, poster) in enumerate(VIDEOS, 1))
    videos = section(row(col(heading("Colaboraciones", "Tres profesionales explican dónde se desbloquea un expediente", "Asesoría, notaría e inmobiliaria: documentación, firma e inmuebles.", "legadia-cases-videos-heading"), "1_1", "legadia-cases-videos-head-col"), "1_1", "legadia-cases-videos-head-row") + row(video_cols, "1_3,1_3,1_3", "legadia-cases-videos-row", make_equal="on", gutter_width="2"), "06 · Tres vídeos", "legadia-cases-section-videos", background_color=COLORS["soft"], custom_padding="96px||96px||true|false")

    results = section(row(col(heading("Qué cambió", "El resultado no es hacer más trámites", "Un caso de éxito se reconoce porque el expediente queda más claro y puede avanzar.", "legadia-cases-results-heading", True), "1_1", "legadia-cases-results-head-col"), "1_1", "legadia-cases-results-head-row") + row(
        col(blurb("Personas localizadas", "Se identifica a quién debe intervenir y cómo contactar.", "legadia-cases-result-1", True, background_color="rgba(255,255,255,0.07)", custom_padding="30px|28px|30px|28px|false|false", border_width_all="1px", border_color_all="rgba(255,255,255,0.20)"), "1_3", "legadia-cases-result-col-1")
        + col(blurb("Documentación ordenada", "Se diferencia lo esencial y se evitan gestiones sin impacto.", "legadia-cases-result-2", True, background_color="rgba(255,255,255,0.07)", custom_padding="30px|28px|30px|28px|false|false", border_width_all="1px", border_color_all="rgba(255,255,255,0.20)"), "1_3", "legadia-cases-result-col-2")
        + col(blurb("Siguiente hito definido", "La familia sabe qué debe ocurrir después y quién lo coordina.", "legadia-cases-result-3", True, background_color="rgba(255,255,255,0.07)", custom_padding="30px|28px|30px|28px|false|false", border_width_all="1px", border_color_all="rgba(255,255,255,0.20)"), "1_3", "legadia-cases-result-col-3"),
        "1_3,1_3,1_3", "legadia-cases-results-row", make_equal="on"), "07 · Resultados que importan", "legadia-cases-section-results", background_color=COLORS["blue"], custom_padding="92px||92px||true|false")

    testimonials = section(row(col(heading("Voces particulares", "Lo que cambia cuando alguien ordena el expediente", "Testimonios anonimizados para proteger la identidad de las familias.", "legadia-cases-testimonials-heading"), "1_1", "legadia-cases-testimonials-head-col"), "1_1", "legadia-cases-testimonials-head-row") + row(
        col(testimonial("Particular de Madrid", "Por primera vez entendimos quién faltaba y qué documentos eran importantes.", 1), "1_3", "legadia-cases-testimonial-col-1")
        + col(testimonial("Familia de Castilla y León", "Dejamos de pedir papeles al azar y todos recibimos la misma explicación.", 2), "1_3", "legadia-cases-testimonial-col-2")
        + col(testimonial("Heredero fuera de España", "Pudimos verificar el motivo del contacto antes de tomar una decisión.", 3), "1_3", "legadia-cases-testimonial-col-3"),
        "1_3,1_3,1_3", "legadia-cases-testimonials-row", make_equal="on"), "08 · Testimonios", "legadia-cases-section-testimonials", background_color=COLORS["white"], custom_padding="92px||92px||true|false")

    cta_copy = heading("Tu caso", "Cuéntanos qué está bloqueado", "Describe la situación, la localidad y qué necesitas resolver. No hace falta tener toda la documentación.", "legadia-cases-cta-heading", True) + button("Llamar ahora", URLS["phone"], "legadia-cases-cta-phone", True) + button("Escribir WhatsApp", URLS["wa"], "legadia-cases-cta-wa")
    cta = section(row(col(cta_copy, "2_5", "legadia-cases-cta-copy-col") + col(contact_form(), "3_5", "legadia-cases-cta-form-col"), "2_5,3_5", "legadia-cases-cta-row", make_equal="on"), "09 · Consulta final", "legadia-cases-section-cta", background_color=COLORS["deep"], background_image=IMAGES["cta"], background_attachment="fixed", use_background_color_gradient="on", background_color_gradient_stops="rgba(3,19,54,0.95) 0%|rgba(8,0,124,0.88) 60%|rgba(224,153,0,0.46) 125%", background_color_gradient_overlays_image="on", custom_padding="100px||100px||true|false")

    footer = section(row(col(blurb("Legadia", "Investigación, tramitación y valoración de derechos hereditarios.", "legadia-cases-footer-brand", True, size="32px"), "1_2", "legadia-cases-footer-brand-col") + col(blurb("Contacto", "91 935 94 72 · tecuidamos@legadia.es · Madrid", "legadia-cases-footer-contact", True, size="24px"), "1_2", "legadia-cases-footer-contact-col"), "1_2,1_2", "legadia-cases-footer-row"), "10 · Footer", "legadia-cases-section-footer", background_color=COLORS["deep"], custom_padding="52px||32px||true|false")

    return header + hero + metrics + featured + carousel + rural + videos + results + testimonials + cta + footer


def build_html() -> str:
    cards = "".join(f'<article><img src="{esc(image_url)}" alt="{esc(place)}"><span>{esc(place)}</span><h3>{esc(title)}</h3><p>{esc(body)}</p><strong>{esc(result)}</strong></article>' for place, title, body, result, image_url in CASES)
    video_cards = "".join(f'<article><div class="poster" style="background-image:linear-gradient(rgba(3,19,54,.15),rgba(3,19,54,.75)),url(\'{esc(poster)}\')"><b>▶</b></div><h3>{esc(title)}</h3><p>{esc(body)}</p></article>' for title, body, poster in VIDEOS)
    return f'''<!doctype html><html lang="es"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Casos de éxito | Legadia</title><style>*{{box-sizing:border-box}}body{{margin:0;font-family:Arial,sans-serif;color:#59647A}}h1,h2,h3{{font-weight:300;color:#0B2B6B}}p{{line-height:1.6}}.wrap{{width:min(1240px,calc(100% - 34px));margin:auto}}header{{background:#0B2B6B;padding:18px;color:#fff}}header .wrap{{display:flex;justify-content:space-between;align-items:center}}header a{{color:#fff;margin-left:18px}}.hero{{padding:110px 0;background:linear-gradient(110deg,rgba(8,0,124,.95),rgba(3,19,54,.92),rgba(224,153,0,.48)),url('{IMAGES["hero"]}') center/cover fixed;color:#fff}}.hero h1{{color:#fff;font-size:64px;max-width:850px}}.hero p{{color:#fff;font-size:20px;max-width:760px}}.metrics{{background:#0B2B6B;color:#fff;padding:38px 0}}.metrics .wrap{{display:grid;grid-template-columns:repeat(4,1fr);text-align:center}}.metrics b{{display:block;color:#E09900;font-size:48px}}section{{padding:88px 0}}.featured{{background:#F4F6FA}}.split{{display:grid;grid-template-columns:1fr 1fr;gap:34px;align-items:center}}.split img{{width:100%;height:520px;object-fit:cover}}.carousel{{background:linear-gradient(rgba(38,28,48,.96),rgba(8,0,124,.91)),url('{IMAGES["rural"]}') center/cover fixed;color:#fff}}.carousel h2,.carousel h3,.carousel p{{color:#fff}}.cards{{display:flex;overflow-x:auto;gap:22px;padding-bottom:18px}}.cards article{{min-width:78%;background:rgba(3,19,54,.88);padding:28px}}.cards img{{width:100%;height:330px;object-fit:cover}}.cards span{{color:#E09900}}.videos{{background:#F4F6FA}}.video-grid,.result-grid,.testimonial-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:22px}}.poster{{height:300px;background-size:cover;background-position:center;display:grid;place-items:center}}.poster b{{background:#fff;color:#0B2B6B;width:76px;height:76px;border-radius:50%;display:grid;place-items:center;font-size:28px}}.results{{background:#0B2B6B;color:#fff}}.results h2,.results h3,.results p{{color:#fff}}.result-grid article{{border:1px solid rgba(255,255,255,.25);padding:26px}}.testimonial-grid article{{border-top:3px solid #E09900;padding:28px;box-shadow:0 14px 40px rgba(11,43,107,.08)}}.cta{{background:linear-gradient(rgba(3,19,54,.95),rgba(8,0,124,.88)),url('{IMAGES["cta"]}') center/cover fixed;color:#fff}}.cta h2,.cta p{{color:#fff}}form{{background:#fff;padding:28px}}input,select,textarea{{width:100%;padding:14px;margin:7px 0;border:0;border-bottom:2px solid #E09900}}textarea{{min-height:120px}}button{{background:#E09900;color:#fff;border:0;padding:15px 24px}}footer{{background:#031336;color:#fff;padding:44px 0}}@media(max-width:850px){{.split,.video-grid,.result-grid,.testimonial-grid,.metrics .wrap{{grid-template-columns:1fr}}.hero h1{{font-size:42px}}header nav{{display:none}}}}</style></head><body><header><div class="wrap"><strong>LEGADIA</strong><nav><a href="{URLS['services']}">Servicios</a><a href="{URLS['cases']}">Casos</a><a href="{URLS['contact']}">Contacto</a></nav></div></header><section class="hero"><div class="wrap"><span>CASOS PARTICULARES</span><h1>Herencias que volvieron a avanzar</h1><p>Casos anonimizados de Madrid, provincias y España rural.</p></div></section><div class="metrics"><div class="wrap"><div><b>+400</b>herederos encontrados</div><div><b>7</b>tipologías</div><div><b>3</b>colaboraciones</div><div><b>1</b>siguiente paso</div></div></div><section class="featured"><div class="wrap split"><img src="{IMAGES['madrid']}" alt="Madrid"><div><span>CASO DESTACADO · MADRID</span><h2>Una vivienda, seis herederos y dos ramas fuera de España</h2><p>La línea sucesoria se reconstruyó y la siguiente fase notarial pudo definirse.</p></div></div></section><section class="carousel"><div class="wrap"><h2>Madrid, provincias y España profunda</h2><div class="cards">{cards}</div></div></section><section><div class="wrap split"><div><h2>Del centro de Madrid a la España rural</h2><p>Registros antiguos, emigración y propiedades de varias generaciones.</p></div><img src="{IMAGES['rural']}" alt="España rural"></div></section><section class="videos"><div class="wrap"><h2>Tres colaboraciones profesionales</h2><div class="video-grid">{video_cards}</div></div></section><section class="results"><div class="wrap"><h2>El resultado no es hacer más trámites</h2><div class="result-grid"><article><h3>Personas localizadas</h3><p>Se identifica a quién debe intervenir.</p></article><article><h3>Documentación ordenada</h3><p>Se evita repetir gestiones.</p></article><article><h3>Siguiente hito definido</h3><p>La familia sabe qué ocurre después.</p></article></div></div></section><section><div class="wrap"><h2>Voces particulares</h2><div class="testimonial-grid"><article>Particular de Madrid: por primera vez entendimos quién faltaba.</article><article>Familia de Castilla y León: dejamos de pedir papeles al azar.</article><article>Heredero fuera de España: verificamos el motivo antes de decidir.</article></div></div></section><section class="cta"><div class="wrap split"><div><h2>Cuéntanos qué está bloqueado</h2><p>No hace falta tener toda la documentación.</p></div><form><input placeholder="Nombre y apellidos"><input placeholder="Teléfono"><input placeholder="Email"><select><option>Localizar herederos</option><option>Ordenar documentación</option><option>No lo tengo claro</option></select><textarea placeholder="Cuéntanos brevemente qué ocurre"></textarea><button>Enviar consulta</button></form></div></section><footer><div class="wrap">Legadia · 91 935 94 72 · tecuidamos@legadia.es · Madrid</div></footer></body></html>'''


def build_docx() -> None:
    doc = Document()
    title = doc.add_heading("Legadia · Casos de éxito", 0)
    title.style.font.name = "Josefin Sans"
    for heading_text in ["Herencias que volvieron a avanzar", "Casos particulares", "Colaboraciones profesionales", "Resultados", "Testimonios", "Consulta final"]:
        doc.add_heading(heading_text, level=1)
        if heading_text == "Casos particulares":
            for place, case_title, body, result, _ in CASES:
                doc.add_heading(f"{place} · {case_title}", level=2)
                doc.add_paragraph(body)
                doc.add_paragraph("Resultado: " + result)
        elif heading_text == "Colaboraciones profesionales":
            for video_title, body, _ in VIDEOS:
                doc.add_heading(video_title, level=2)
                doc.add_paragraph(body)
    styles = doc.styles
    styles["Normal"].font.name = "Josefin Sans"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor(89, 100, 122)
    doc.save(OUT / "Legadia-Casos-de-Exito-Textos-Exactos.docx")


def main() -> None:
    layout = build_layout()
    payload = {"context": "et_builder", "data": {"900305": layout}, "presets": [], "global_colors": []}
    json_text = json.dumps(payload, ensure_ascii=False, indent=2)
    html_text = build_html()
    bad = ["Ã", "Â", "â", "�", "\\u003c", "\\u003e", "&lt;", "&gt;"]
    if any(pattern in json_text or pattern in html_text for pattern in bad):
        raise RuntimeError("Mojibake o HTML escapado visible detectado")
    if "[et_pb_text" in layout or "[et_pb_code" in layout:
        raise RuntimeError("Se detectó un módulo prohibido")
    modules = re.findall(r"\[(et_pb_[a-z_]+)\s+([^\]]*)\]", layout)
    missing_classes = [name for name, attr_text in modules if "module_class=" not in attr_text]
    if missing_classes:
        raise RuntimeError("Módulos sin clase CSS: " + ", ".join(missing_classes[:10]))
    json_path = OUT / "Legadia-Casos-de-Exito-Enfocada-Divi4.json"
    html_path = OUT / "Legadia-Casos-de-Exito-Enfocada-Preview.html"
    validation_path = OUT / "VALIDACION-Casos-de-Exito-Enfocada.txt"
    json_path.write_text(json_text, encoding="utf-8")
    html_path.write_text(html_text, encoding="utf-8")
    build_docx()
    validation_path.write_text(
        "VALIDACIÓN CASOS DE ÉXITO ENFOCADA · DIVI 4\nEstado: CORRECTO\nSecciones: 11\nMódulos Slider: 1\nSlides de casos: 7\nMódulos Vídeo: 3\nTestimonios: 3\nFormularios: 1\nMódulos Texto: 0\nMódulos Código: 0\nTodos los módulos incluyen clase CSS editable.\nSin mojibake ni HTML escapado visible.\n",
        encoding="utf-8",
    )
    zip_path = OUT / "Legadia-Casos-de-Exito-Enfocada-Divi4.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in [json_path, html_path, OUT / "Legadia-Casos-de-Exito-Textos-Exactos.docx", validation_path]:
            zf.write(path, path.name)


if __name__ == "__main__":
    main()
